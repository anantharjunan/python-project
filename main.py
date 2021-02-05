from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


if __name__ == '__main__':
    document = Document()
    # profile picture
    document.add_picture('profile.png', width=Inches(1.0))

    # name, phone number and email details
    name = input('What is your name? ')
    speak('Hello ' + name + ' how are you today?')
    speak('What is your phone number?')
    phone_number = input('What is your phone number? ')
    speak('What is your email?')
    email = input('What is your email? ')
    document.add_paragraph(
        name + ' | ' + phone_number + ' | ' + email
    )

    # about me
    document.add_heading('About me')
    speak('Tell about yourself?')
    about_me = input('Tell about yourself? ')
    document.add_paragraph(about_me)

    # Experience
    document.add_heading('Experience')
    p = document.add_paragraph()
    speak('Can you name the club?')
    company = input('Enter club ')
    from_date = input('From Date ')
    to_date = input('To Date ')

    p.add_run(company + ' ').bold = True
    p.add_run(from_date + '-' + to_date + '\n').italic = True
    speak('Describe your experience at ' + company + '?')
    experience_details = input('Describe your experience at ' + company + ' ')
    p.add_run(experience_details)

    # more experiences
    while True:
        speak('Do you have more experience? Yes or No')
        has_more_experiences = input('Do you have more experience? Yes or No ')
        if has_more_experiences.lower() == 'yes':
            p = document.add_paragraph()
            speak('Can you name the club?')
            company = input('Enter club ')
            from_date = input('From Date ')
            to_date = input('To Date ')

            p.add_run(company + ' ').bold = True
            p.add_run(from_date + '-' + to_date + '\n').italic = True
            speak('Describe your experience at ' + company + '?')
            experience_details = input('Describe your experience at ' + company + ' ')
            p.add_run(experience_details)
        else:
            break

    # Skills
    document.add_heading('Skills')
    speak('Can you describe your skill')
    skill = input('Enter skill ')
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'

    while True:
        speak('Do you have more skills? Yes or No')
        has_more_stills = input('Do you have more skills? Yes or No ')
        if has_more_stills.lower() == 'yes':
            speak('Can you describe your skill')
            skill = input('Enter skill ')
            p = document.add_paragraph(skill)
            p.style = 'List Bullet'
        else:
            break

    document.save('profile.docx')
